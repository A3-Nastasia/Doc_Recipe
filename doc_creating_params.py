from docx import Document
import os
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocRecipe:
    def __init__(self, doc_name):
        """
        DESCRIPTION: initialize class object

        Args:
            self (DocRecipe): Instance of the current class
            doc_name (str): name of the .docx file
        """
        # Create a doc object from python-docx
        self.doc = Document()
        self.doc_name = doc_name

    # Saving file and checking if it's opened
    def __save_docx(self, doc):
        """
        DESCRIPTION: Save file and check if it's opened

        Args:
            self (DocRecipe): Instance of the current class
            doc (Document): Document object from python-docx
        """
        try:
            doc_path = os.path.join(os.getcwd(), self.doc_name)
            doc.save(self.doc_name)
            print("File saved on path", doc_path)
        except PermissionError:
            print(f"Error: The file {self.doc_name} is currently in use. Please close it and try again.")

    def __set_doc_format(self,
                         doc,
                         font_size = 14,
                         doc_margin_top_in_cm = 1,
                         doc_margin_right_in_cm = 1,
                         doc_margin_left_in_cm = 1,
                         doc_margin_bottom_in_cm = 1,
                         paragraphs_space_after = 0):
        """
        DESCRIPTION: Set format for .docx text

        Args:
            self (DocRecipe): Instance of the current class
            doc (Document): Document object from python-docx

            font_size (int): Font size of all text in document

            doc_margin_top_in_cm (int): Document's margin of top
            doc_margin_right_in_cm (int): Document's margin of right
            doc_margin_left_in_cm (int): Document's margin of left
            doc_margin_bottom_in_cm (int): Document's margin of bottom

            paragraphs_space_after (int): Space after paragraphs in document
        """
        # Set font size as default for all doc
        style = doc.styles["Normal"]
        style.font.size = Pt(font_size)

        # Set margin for sheet(-s)
        section = doc.sections[0]
        section.top_margin = Cm(doc_margin_top_in_cm)
        section.right_margin = Cm(doc_margin_right_in_cm)
        section.left_margin = Cm(doc_margin_left_in_cm)
        section.bottom_margin = Cm(doc_margin_bottom_in_cm)

        # Set none paragraph space after
        style.paragraph_format.space_after = paragraphs_space_after

    def __set_table_col_width(self, table, col_id, width_in_cm):
        """
        DESCRIPTION: Set column width in Cm

        Args:
            self (DocRecipe): Instance of the current class
            table (Table): Table object inside .docx (Document)

            col_id (int): Id of column to set width
            width_in_cm (int): Width of column in centimetres
        """
        for cell in table.columns[col_id].cells:
            cell.width = Cm(width_in_cm)

    def __set_recipe_name_style(self, recipe_name, font_size, space_after):
        """
        DESCRIPTION: Set recipe name and style for it

        Args:
            self (DocRecipe): Instance of the current class
            recipe_name (str): Recipe name
            font_size (int): Font size in Pt
            space_after (int): Space after paragraph in Pt
        """
        recipe_name_heading = self.doc.add_paragraph()
        run_heading = recipe_name_heading.add_run(recipe_name)
        run_heading.font.size = Pt(font_size)
        run_heading.font.italic = True
        run_heading.font.bold = True
        recipe_name_heading.paragraph_format.space_after = Pt(space_after)

    def create_doc(self,
                   recipe_name,
                   table_grid=False):
        """
        DESCRIPTION: Create a .docx file with recipe name and save it

        Args:
            self (DocRecipe): Instance of the current class
            recipe_name (str): Recipe name
            table_grid (bool): Set table grid
        """
        self.__set_doc_format(self.doc)
        self.__set_recipe_name_style(recipe_name, 16, 20)

        table = self.doc.add_table(rows=1, cols=3)
        # If the variable is set
        if table_grid:
            table.style = 'Table Grid'

        # Do not stretch the table (text within)
        table.autofit = False

        # Header of the table
        hdr_cells = table.rows[0].cells
        table.cell(0,0).merge(table.cell(0,1))
        hdr_cells[0].paragraphs[0].add_run("Ingredients").font.bold = True
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].paragraphs[0].add_run("Cooking steps").font.bold = True
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        amount_ingredients = 5
        measure_ingredient = "kg" # Change to list
        for i in range(amount_ingredients):
            row = table.add_row()
            row.cells[0].text = f"{{{i}#_amount_of_this_ingredient}}" + "\n" + f"{measure_ingredient}"
            row.cells[1].text = "{name_of_ingredient}"
            row.cells[2].text = "{Cooking_steps_of_recipe _and_something_to_text_here}"

        # Merge (connect) all cols at the right side
        # The list of ingredients is at the left side
        # To write a proper version of cooking steps it should be one space
        table.cell(1,2).merge(table.cell(amount_ingredients,2))

        self.__set_table_col_width(table, 0, 2)
        self.__set_table_col_width(table, 1, 4.5)
        self.__set_table_col_width(table, 2, 13.5)

        self.__save_docx(self.doc)


test_doc = DocRecipe("Recipe_template.docx")
recipe_name = "Recipe_name"
test_doc.create_doc(recipe_name,
                    table_grid=False)